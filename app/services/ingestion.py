"""Knowledge-base ingestion into the local Qdrant vector store (hybrid RAG).

PDF/DOCX/EPUB/TXT/MD files are normalised to Markdown, chunked, embedded and
upserted *locally*: ``to_markdown`` converts via markitdown (falling back to
``pypdf`` / ``python-docx`` / a tiny EPUB parser), the text is split into
overlapping windows, embedded by the bge-m3 sidecar (dense + sparse) and written
to Qdrant. There is no hosted vector store — the KB lives in the Qdrant
collection.

The full ingest path is::

    upload_document
      -> to_markdown    (normalise pdf/docx/epub/txt/md -> Markdown)
      -> _chunk_text    (sliding character windows, CHUNK_SIZE/CHUNK_OVERLAP)
      -> embeddings.embed_texts   (bge-m3 dense + sparse)
      -> vectorstore.upsert_points   (payload carries lab metadata)

``bulk_ingest_tree`` walks the school corpus folder, derives per-file metadata
from the path (``corpus_meta.parse_path`` -> subject/grade/lang/lab_id) and
ingests each file with that metadata so retrieval can be scoped by lab context.
The document id is a stable ``uuid5`` of the ``doc_key`` (the relative path for
bulk ingest, since lab filenames repeat across grades). Re-ingest stages and
activates a new generation before stale chunks are removed.

Validation failures (unsupported extension, etc.) raise ``ValueError``; any
embedding/vector-store failure surfaces as an ``LLMError`` subclass and is left
to propagate so routes can map it to 504/502.
"""

import asyncio
import io
import json
import logging
import posixpath
import re
import uuid
import xml.etree.ElementTree as ET
import zipfile
from collections.abc import Awaitable, Callable
from collections import Counter
from html import unescape as _html_unescape
from pathlib import Path
from urllib.parse import unquote

import docx
import pypdf

from app.core.config import settings
from app.services import corpus_meta, embeddings, vectorstore

logger = logging.getLogger("assistant.ingestion")

# File types we can parse and embed locally for the school knowledge base.
# Everything is normalised to Markdown before chunking (see ``to_markdown``).
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".epub"}

# uuid5 namespace for deriving stable document/chunk ids from doc keys.
_DOC_NAMESPACE = uuid.UUID("6ba7b811-9dad-11d1-80b4-00c04fd430c8")

# Documents shorter than this many characters of extracted text are flagged as
# "stub" in the manifest (present but too thin to be a real lab procedure).
_STUB_CHARS = 200

_TAG_RE = re.compile(r"<[^>]+>")
# Drop <script>/<style> bodies before stripping tags so their contents never
# leak into the extracted text.
_SCRIPT_STYLE_RE = re.compile(r"(?is)<(script|style)\b[^>]*>.*?</\1>")
# A word for image-only and extraction-quality heuristics: two or more Unicode
# letters, excluding digits and underscores. This recognizes English, Russian,
# and Kazakh prose while still ignoring page numbers and markup identifiers.
_WORD_RE = re.compile(r"[^\W\d_]{2,}", re.UNICODE)

# Tokens and known artifacts used to assess PDF text-layer quality. Some
# chemistry PDFs have a broken text layer made almost entirely of the OKULYK
# redistribution notice and synthetic page labels. That text has plenty of
# Cyrillic words, so word count alone cannot distinguish it from book content.
_INFO_TOKEN_RE = re.compile(r"[^\W_]{2,}", re.UNICODE)
_OKULYK_RE = re.compile(r"(?i)\bokulyk\.(?:com|kz)\b")
_PAGE_TOKEN_RE = re.compile(r"(?i)(?<!\w)page\s*\d{1,5}(?!\w)")
_MARKDOWN_RULE_RE = re.compile(r"^[\s|:+\-_*`#>]+$")
_KNOWN_BOILERPLATE_PREFIXES = (
    "книга предоставлена исключительно в образовательных целях",
    "согласно приказа министра образования и науки республики казахстан",
    "приказа министра образования и науки республики казахстан",
    "республики казахстан от 17 мая 2019 года",
    "от 17 мая 2019 года № 217",
    "все учебники казахстана ищите на сайтах",
    "все учебники казахстана на",
)

# Headings worth carrying into citation metadata. Markdown headings are kept
# verbatim by markitdown; the plain-text alternatives cover the most common
# chapter/section labels in the Russian, Kazakh, and English school corpus.
_HEADING_RE = re.compile(
    r"(?im)^\s*(?:(?P<markdown>#{1,6})\s+)?"
    r"(?P<title>(?:(?:глава|раздел|параграф|тарау|бөлім|chapter|section|unit|lesson)\b|\xa7)"
    r"[^\r\n]{0,150})\s*$"
)

# EPUB extraction thresholds (word counts via ``_WORD_RE``):
#   * if markitdown returns at least this many words we trust it and skip the
#     (more expensive) spine reparse;
#   * below ``_EPUB_MIN_WORDS`` the best extraction is treated as image-only
#     (scanned) and the file is skipped rather than indexed as a few noise
#     chunks of stripped <img> tags.
_EPUB_TRUST_WORDS = 500
_EPUB_MIN_WORDS = 50

# OCR (opt-in, ingest-time only). Image members inside an EPUB and the language
# mapping for Tesseract. Imports of pytesseract / pypdfium2 / PIL are deferred to
# the OCR helpers so the serving path (and the test suite) never need them.
_IMG_SRC_RE = re.compile(r"""<img\b[^>]*?\bsrc\s*=\s*['"]([^'"]+)['"]""", re.IGNORECASE)
_IMG_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".webp")

ProgressCallback = Callable[[str], Awaitable[None]]
CancelCheck = Callable[[], Awaitable[bool]]


class IngestionCancelled(Exception):
    """Raised when an administrator cancels ingestion at a safe boundary."""


def _tesseract_lang(lang: str | None) -> str:
    """Map the corpus ``lang`` code to Tesseract model name(s)."""
    return {"ru": "rus", "kk": "kaz", "en": "eng"}.get(
        lang or "", "rus+kaz+eng"
    )


def _ocr_image(image, lang: str) -> str:
    """OCR a single PIL image with Tesseract (deferred import)."""
    import pytesseract

    return pytesseract.image_to_string(image, lang=lang)


def _ocr_pdf(content: bytes, lang: str) -> str:
    """Render a (scanned) PDF with pypdfium2 and OCR each page.

    Pure-wheel rendering (no poppler): page -> bitmap -> PIL -> Tesseract.
    """
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(content)
    try:
        n_pages = len(pdf)
        max_pages = settings.OCR_MAX_PAGES or n_pages
        scale = settings.OCR_DPI / 72.0  # pdfium renders at 72 dpi * scale
        out: list[str] = []
        for i in range(min(n_pages, max_pages)):
            page = pdf[i]
            bitmap = page.render(scale=scale)
            image = bitmap.to_pil()
            out.append(_ocr_image(image, lang))
        return "\n".join(out)
    finally:
        pdf.close()


def _epub_image_members(z: zipfile.ZipFile, names: list[str]) -> list[str]:
    """EPUB image members in spine/reading order (fallback: all images, sorted).

    Follows the spine documents (reusing :func:`_spine_docs` / :func:`_zip_join`)
    and collects each ``<img src=…>`` member in reading order, de-duplicated.
    When no spine images resolve, falls back to every image member by name.
    """
    member_set = set(names)
    opf_path = _opf_path(z, names)
    docs = _spine_docs(z, opf_path) if opf_path else []
    docs = [d for d in docs if d in member_set]

    ordered: list[str] = []
    seen: set[str] = set()
    for doc in docs:
        base = doc.rsplit("/", 1)[0] if "/" in doc else ""
        try:
            html = z.read(doc).decode("utf-8", "ignore")
        except KeyError:
            continue
        for m in _IMG_SRC_RE.finditer(html):
            member = _zip_join(base, m.group(1))
            if member in member_set and member not in seen:
                seen.add(member)
                ordered.append(member)

    if not ordered:
        ordered = sorted(n for n in names if n.lower().endswith(_IMG_EXTS))
    return ordered


def _ocr_epub_images(content: bytes, lang: str) -> str:
    """OCR the page-images of a scanned EPUB, in reading order."""
    from PIL import Image

    with zipfile.ZipFile(io.BytesIO(content)) as z:
        members = _epub_image_members(z, z.namelist())
        max_pages = settings.OCR_MAX_PAGES or len(members)
        out: list[str] = []
        for name in members[:max_pages]:
            try:
                data = z.read(name)
            except KeyError:
                continue
            try:
                image = Image.open(io.BytesIO(data))
                out.append(_ocr_image(image, lang))
            except Exception as exc:  # noqa: BLE001 - skip an undecodable image
                logger.warning("OCR skipped EPUB image '%s' (%s)", name, exc)
        return "\n".join(out)


def _doc_id(doc_key: str) -> str:
    """Stable document id for ``doc_key`` (re-ingest replaces in place).

    ``doc_key`` is the bare filename for single uploads but the *relative path*
    for bulk corpus ingest, since lab filenames (e.g. 'Лабораторная работа
    №2.docx') repeat across subjects/grades and would otherwise collide.
    """
    return uuid.uuid5(_DOC_NAMESPACE, doc_key).hex


def _markitdown(suffix: str, content: bytes) -> str | None:
    """Convert ``content`` to Markdown via markitdown, or None if unavailable.

    markitdown (optional dep) handles pdf/docx/epub/xlsx richly; when it is not
    installed or fails we fall back to the per-format extractors below.
    """
    try:
        from markitdown import MarkItDown
    except ImportError:
        return None
    try:
        result = MarkItDown().convert_stream(io.BytesIO(content), file_extension=suffix)
        text = (result.text_content or "").strip()
        return text or None
    except Exception as exc:  # noqa: BLE001 - any converter failure -> fallback
        logger.warning("markitdown failed for '%s' (%s); using fallback", suffix, exc)
        return None


def _count_words(text: str) -> int:
    """Number of word-ish tokens in ``text`` (cheap image-only heuristic)."""
    return len(_WORD_RE.findall(text))


def _information_tokens(text: str) -> list[str]:
    """Return case-folded word/number tokens for extraction-quality checks."""
    return [token.casefold() for token in _INFO_TOKEN_RE.findall(text)]


def _is_pdf_artifact_line(line: str) -> bool:
    """Whether ``line`` is a known watermark, page label or table artifact."""
    stripped = line.strip()
    if not stripped:
        return False

    folded = re.sub(r"\s+", " ", stripped).casefold().strip(" *_`#>|-")
    if _OKULYK_RE.search(folded):
        return True
    if any(folded.startswith(prefix) for prefix in _KNOWN_BOILERPLATE_PREFIXES):
        return True
    if _MARKDOWN_RULE_RE.fullmatch(stripped) and any(ch in stripped for ch in "|-_"):
        return True

    # Covers `page64`, `| page65 | | 65 |`, and long rows/sequences made only
    # from page numbers and Markdown table punctuation. A lone number is kept
    # because it may be an exercise number rather than a page marker.
    plain = re.sub(r"[|*_`#>:()\[\]{}+\-/\\]", " ", stripped)
    plain = re.sub(r"\s+", " ", plain).strip()
    if re.fullmatch(r"(?i)page\s*\d{1,5}(?:\s+\d{1,5})*", plain):
        return True
    numeric_parts = re.findall(r"\d{1,5}", plain)
    without_numbers = re.sub(r"\d{1,5}", "", plain).strip()
    return not without_numbers and len(numeric_parts) >= 4


def _clean_pdf_extraction(text: str) -> str:
    """Remove known PDF text-layer boilerplate without rewriting book prose.

    Explicit OKULYK/legal/page artifacts are always removed. Unknown repeated
    lines are removed only when short, non-sentence lines overwhelmingly
    dominate the extraction. This deliberately keeps repeated definitions,
    worked examples and other normal textbook paragraphs.
    """
    if not text:
        return ""

    lines: list[str] = []
    for raw_line in text.splitlines():
        if _is_pdf_artifact_line(raw_line):
            continue
        # A page token can occasionally be prefixed/suffixed to useful text on
        # the same extracted line. Remove that token but retain the real text.
        line = _PAGE_TOKEN_RE.sub(" ", raw_line)
        line = re.sub(r"[ \t]+", " ", line).strip()
        if line:
            lines.append(line)

    if not lines:
        return ""

    keys = [re.sub(r"\s+", " ", line).casefold() for line in lines]
    counts = Counter(keys)
    token_counts = {
        key: len(_information_tokens(line))
        for key, line in zip(keys, lines, strict=True)
    }
    total_tokens = sum(token_counts[key] for key in keys)
    repeated_keys = {
        key
        for key, count in counts.items()
        if count >= 4
        and 0 < token_counts[key] <= 12
        and len(key) <= 140
        and not re.search(r"[.!?…][\s*_`]*$", key)
    }
    repeated_tokens = sum(counts[key] * token_counts[key] for key in repeated_keys)
    repeated_vocabulary = {
        token for key in repeated_keys for token in _information_tokens(key)
    }
    if (
        repeated_keys
        and total_tokens
        and repeated_tokens / total_tokens >= 0.75
        and len(repeated_vocabulary) <= 24
    ):
        lines = [line for line, key in zip(lines, keys, strict=True) if key not in repeated_keys]

    return "\n".join(lines).strip()


def _is_low_quality_pdf_extraction(raw_text: str, cleaned_text: str) -> bool:
    """Detect thin or overwhelmingly repetitive PDF text-layer extraction."""
    clean_words = _count_words(cleaned_text)
    if clean_words < _EPUB_MIN_WORDS:
        return True

    raw_tokens = _information_tokens(raw_text)
    clean_tokens = _information_tokens(cleaned_text)
    if len(raw_tokens) >= 100 and len(clean_tokens) / len(raw_tokens) < 0.35:
        return True

    # Catch unknown watermark phrases repeated hundreds of times even when they
    # contain enough Cyrillic words to pass the old threshold. Thresholds are
    # intentionally strict so repeated textbook paragraphs remain usable.
    if len(clean_tokens) >= 100:
        lexical_diversity = len(set(clean_tokens)) / len(clean_tokens)
        window_count = len(clean_tokens) - 5
        unique_windows = {
            tuple(clean_tokens[i : i + 6]) for i in range(window_count)
        }
        window_diversity = len(unique_windows) / window_count
        if lexical_diversity < 0.035 and window_diversity < 0.08:
            return True
    return False


def _materially_better_pdf_text(
    raw_candidate: str,
    cleaned_candidate: str,
    raw_baseline: str,
    cleaned_baseline: str,
) -> bool:
    """Whether cleaned OCR output is a usable improvement over the text layer."""
    candidate_low_quality = _is_low_quality_pdf_extraction(
        raw_candidate, cleaned_candidate
    )
    baseline_low_quality = _is_low_quality_pdf_extraction(raw_baseline, cleaned_baseline)
    if baseline_low_quality and not candidate_low_quality:
        return True
    if candidate_low_quality:
        return False

    candidate_tokens = _information_tokens(cleaned_candidate)
    baseline_tokens = _information_tokens(cleaned_baseline)
    return len(candidate_tokens) >= max(
        len(baseline_tokens) + 25, len(baseline_tokens) * 1.2
    )


def _html_to_text(html: str) -> str:
    """Strip an (x)html document to readable text (drops script/style, entities)."""
    html = _SCRIPT_STYLE_RE.sub(" ", html)
    text = _TAG_RE.sub(" ", html)
    return _html_unescape(text)


def _zip_join(base: str, href: str) -> str:
    """Resolve an OPF-relative href to a normalised zip member path."""
    href = unquote(href).split("#", 1)[0]  # drop url-encoding + fragment
    joined = f"{base}/{href}" if base else href
    return posixpath.normpath(joined)


def _opf_path(z: zipfile.ZipFile, names: list[str]) -> str | None:
    """Locate the OPF package document via META-INF/container.xml (or first .opf)."""
    if "META-INF/container.xml" in names:
        try:
            root = ET.fromstring(z.read("META-INF/container.xml"))
            for el in root.iter():
                if el.tag.rsplit("}", 1)[-1] == "rootfile":
                    full_path = el.get("full-path")
                    if full_path:
                        return full_path
        except ET.ParseError:
            pass
    for name in names:
        if name.lower().endswith(".opf"):
            return name
    return None


def _spine_docs(z: zipfile.ZipFile, opf_path: str) -> list[str]:
    """Return spine document zip paths in reading order, resolved against the OPF."""
    try:
        opf = ET.fromstring(z.read(opf_path))
    except (KeyError, ET.ParseError):
        return []
    manifest: dict[str, str] = {}
    for el in opf.iter():
        if el.tag.rsplit("}", 1)[-1] == "item":
            item_id, href = el.get("id"), el.get("href")
            if item_id and href:
                manifest[item_id] = href
    base = opf_path.rsplit("/", 1)[0] if "/" in opf_path else ""
    docs: list[str] = []
    for el in opf.iter():
        if el.tag.rsplit("}", 1)[-1] == "itemref":
            href = manifest.get(el.get("idref") or "")
            if href:
                docs.append(_zip_join(base, href))
    return docs


def _epub_to_text(content: bytes) -> str:
    """Robust EPUB -> text: follow the OPF spine, in reading order.

    markitdown (and ebooklib) sometimes return almost nothing for these school
    textbooks because the spine isn't followed, so we parse the OCF container
    ourselves: META-INF/container.xml -> OPF manifest+spine -> each spine
    document, tags stripped, concatenated in order. If the container/spine can't
    be read we fall back to every (x)html member, sorted alphabetically (the old
    behaviour).
    """
    with zipfile.ZipFile(io.BytesIO(content)) as z:
        names = z.namelist()
        member_set = set(names)
        opf_path = _opf_path(z, names)
        docs = _spine_docs(z, opf_path) if opf_path else []
        docs = [d for d in docs if d in member_set]
        if not docs:
            docs = sorted(
                n for n in names if n.lower().endswith((".xhtml", ".html", ".htm"))
            )
        out: list[str] = []
        for name in docs:
            try:
                raw = z.read(name)
            except KeyError:
                continue
            out.append(_html_to_text(raw.decode("utf-8", "ignore")))
    return "\n".join(out)


def _extract_epub(
    filename: str, content: bytes, *, ocr: bool = False, lang: str | None = None
) -> str:
    """EPUB -> text, resilient to markitdown silently extracting nothing.

    markitdown is tried first; when it returns suspiciously little text we
    re-parse the spine ourselves (:func:`_epub_to_text`) and keep whichever
    yields more words. Genuinely image-only (scanned) EPUBs extract ~no words:
    when ``ocr`` is set we OCR the page-images instead of skipping; otherwise we
    log a clear warning naming the file and return ``""`` so the document is
    skipped (status ``empty``) rather than indexed as a few noise chunks. Those
    files are listed in the manifest as ``stub`` for the OCR decision.
    """
    md = _markitdown(".epub", content) or ""
    md_words = _count_words(md)

    best, best_words = md, md_words
    if md_words < _EPUB_TRUST_WORDS:
        try:
            spine = _epub_to_text(content)
        except Exception as exc:  # noqa: BLE001 - any zip/xml failure -> keep markitdown
            logger.warning("EPUB spine parse failed for '%s' (%s)", filename, exc)
            spine = ""
        if _count_words(spine) > best_words:
            best, best_words = spine, _count_words(spine)

    if best_words < _EPUB_MIN_WORDS:
        if ocr:
            tlang = _tesseract_lang(lang)
            logger.info(
                "EPUB '%s' is image-only (%d words); running OCR (lang=%s)",
                filename, best_words, tlang,
            )
            ocr_text = _ocr_epub_images(content, tlang)
            ocr_words = _count_words(ocr_text)
            logger.info("OCR recovered %d words from EPUB '%s'", ocr_words, filename)
            if ocr_words >= _EPUB_MIN_WORDS:
                return ocr_text
        logger.warning(
            "EPUB '%s' yielded ~no extractable text (%d words) — likely "
            "image-only/scanned; skipping (needs OCR)",
            filename,
            best_words,
        )
        return ""
    return best


def to_markdown(
    filename: str, content: bytes, *, ocr: bool = False, lang: str | None = None
) -> str:
    """Normalise any supported document to Markdown / plain text.

    Tries markitdown first (best fidelity, incl. tables and EPUB); otherwise
    falls back to pypdf / python-docx / a tiny EPUB tag-stripper. ``.md`` /
    ``.txt`` pass through unchanged. Raises ``ValueError`` for unsupported
    extensions.

    PDF text is cleaned of known redistribution/page boilerplate before it is
    returned. When ``ocr`` is set and a scanned EPUB/PDF extraction is thin or
    overwhelmingly repetitive, the page images are OCR'd (Tesseract,
    ``lang``-aware). OCR is opt-in and ingest-time only.
    """
    suffix = Path(filename).suffix.lower()

    if suffix in (".txt", ".md"):
        return content.decode("utf-8", errors="replace")

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Supported: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    parsed_docx = None
    if suffix == ".pdf":
        if b"%PDF-" not in content[:1024]:
            raise ValueError("Invalid PDF file")
        try:
            reader = pypdf.PdfReader(io.BytesIO(content))
            len(reader.pages)
        except Exception as exc:  # noqa: BLE001 - normalise parser failures
            raise ValueError("Invalid PDF file") from exc
    if suffix in {".docx", ".epub"}:
        label = suffix.removeprefix(".").upper()
        try:
            if not zipfile.is_zipfile(io.BytesIO(content)):
                raise ValueError
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = set(archive.namelist())
                if suffix == ".docx" and not {
                    "[Content_Types].xml",
                    "word/document.xml",
                }.issubset(names):
                    raise ValueError
                if suffix == ".epub":
                    mimetype = archive.read("mimetype").decode("ascii").strip()
                    if mimetype != "application/epub+zip":
                        raise ValueError
                    if "META-INF/container.xml" in names:
                        container = ET.fromstring(
                            archive.read("META-INF/container.xml")
                        )
                        if container.tag.rsplit("}", 1)[-1] != "container":
                            raise ValueError
                        opf_path = next(
                            element.get("full-path")
                            for element in container.iter()
                            if element.tag.rsplit("}", 1)[-1] == "rootfile"
                            and element.get("full-path")
                        )
                        if opf_path not in names:
                            raise ValueError
                        package = ET.fromstring(archive.read(opf_path))
                        if package.tag.rsplit("}", 1)[-1] != "package":
                            raise ValueError
                    elif not any(
                        name.lower().endswith(
                            (".xhtml", ".html", ".htm", *_IMG_EXTS)
                        )
                        for name in names
                    ):
                        raise ValueError
            if suffix == ".docx":
                parsed_docx = docx.Document(io.BytesIO(content))
        except Exception as exc:  # noqa: BLE001 - normalise parser failures
            raise ValueError(f"Invalid {label} file") from exc

    # EPUB has its own markitdown-first-then-spine-fallback path (markitdown
    # silently under-reads the spine on these textbooks), so handle it before the
    # generic markitdown attempt below.
    if suffix == ".epub":
        return _extract_epub(filename, content, ocr=ocr, lang=lang)

    md = _markitdown(suffix, content)

    if suffix == ".pdf":
        if md is not None:
            raw_text = md
        else:
            reader = pypdf.PdfReader(io.BytesIO(content))
            raw_text = "\n".join(p.extract_text() or "" for p in reader.pages)
        text = _clean_pdf_extraction(raw_text)
        low_quality = _is_low_quality_pdf_extraction(raw_text, text)
        # OCR thin or corrupt text layers when asked, then compare the cleaned
        # candidates so a larger watermark-heavy OCR result cannot win.
        if ocr and low_quality:
            tlang = _tesseract_lang(lang)
            logger.info(
                "PDF '%s' has low-quality text (%d raw/%d cleaned words); "
                "running OCR (lang=%s)",
                filename,
                _count_words(raw_text),
                _count_words(text),
                tlang,
            )
            raw_ocr_text = _ocr_pdf(content, tlang)
            ocr_text = _clean_pdf_extraction(raw_ocr_text)
            logger.info(
                "OCR recovered %d cleaned words from PDF '%s'",
                _count_words(ocr_text),
                filename,
            )
            if _materially_better_pdf_text(raw_ocr_text, ocr_text, raw_text, text):
                return ocr_text
        return text

    if md is not None:
        return md

    if suffix == ".docx":
        return "\n".join(
            p.text for p in parsed_docx.paragraphs if p.text and p.text.strip()
        )

    raise ValueError(f"Unsupported file type '{suffix}'")  # pragma: no cover


# Backwards-compatible alias (older callers / tests imported ``_extract_text``).
_extract_text = to_markdown


def _chunk_text(text: str) -> list[str]:
    """Split ``text`` into overlapping character windows.

    Whitespace is collapsed first, then windows of ``CHUNK_SIZE`` characters are
    emitted stepping by ``CHUNK_SIZE - CHUNK_OVERLAP``. Empty / whitespace-only
    windows are dropped.
    """
    normalized = _normalize_text(text)
    if not normalized:
        return []

    size = settings.CHUNK_SIZE
    overlap = settings.CHUNK_OVERLAP
    # Guard against a degenerate (or misconfigured) overlap >= size which would
    # make the window never advance.
    if overlap >= size:
        overlap = max(0, size // 4)
    step = size - overlap

    chunks: list[str] = []
    for start in range(0, len(normalized), step):
        window = normalized[start : start + size].strip()
        if window:
            chunks.append(window)
        if start + size >= len(normalized):
            break
    return chunks


def _normalize_text(text: str) -> str:
    """Collapse whitespace exactly as the chunker does."""
    return re.sub(r"\s+", " ", text).strip()


def _section_markers(text: str) -> list[tuple[int, str, str]]:
    """Return normalized offsets for chapter/section headings in ``text``.

    The offsets are approximate only at whitespace boundaries, which is enough
    for assigning the latest heading to a character-window chunk. A heading is
    metadata, not part of retrieval scoring or document reconstruction.
    """
    markers: list[tuple[int, str, str]] = []
    for match in _HEADING_RE.finditer(text):
        title = _normalize_text(match.group("title"))
        if not title:
            continue
        prefix = _normalize_text(text[: match.start()])
        offset = len(prefix) + (1 if prefix else 0)
        kind = (
            "chapter"
            if title.casefold().startswith(("глава", "тарау", "chapter", "unit"))
            else "section"
        )
        markers.append((offset, kind, title))
    return markers


def _pdf_page_spans(content: bytes, extracted_text: str) -> list[tuple[int, int, int]]:
    """Map normalized PDF text offsets to 1-based page numbers when safe.

    Page metadata is emitted only when pypdf's page-by-page extraction exactly
    matches the normalized text selected by :func:`to_markdown`. This avoids
    attaching incorrect pages when markitdown produced materially different
    text, or when OCR was needed. The serving and OCR paths remain unchanged.
    """
    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = [
            (page_number, _normalize_text(page.extract_text() or ""))
            for page_number, page in enumerate(reader.pages, start=1)
        ]
    except Exception as exc:  # noqa: BLE001 - citation locator is best-effort
        logger.debug("Could not derive PDF page locators: %s", exc)
        return []

    nonempty = [(number, text) for number, text in pages if text]
    normalized = _normalize_text(extracted_text)
    if not nonempty or " ".join(text for _, text in nonempty) != normalized:
        return []

    spans: list[tuple[int, int, int]] = []
    cursor = 0
    for page_number, page_text in nonempty:
        start = cursor
        end = start + len(page_text)
        spans.append((start, end, page_number))
        cursor = end + 1
    return spans


def _chunk_records(
    text: str, *, page_spans: list[tuple[int, int, int]] | None = None
) -> list[dict]:
    """Chunk text and retain character, page and heading locators per chunk."""
    normalized = _normalize_text(text)
    if not normalized:
        return []

    size = settings.CHUNK_SIZE
    overlap = settings.CHUNK_OVERLAP
    if overlap >= size:
        overlap = max(0, size // 4)
    step = size - overlap
    markers = _section_markers(text)

    records: list[dict] = []
    active: dict[str, str] = {}
    marker_index = 0
    for start in range(0, len(normalized), step):
        end = min(start + size, len(normalized))
        raw_window = normalized[start:end]
        window = raw_window.strip()
        if not window:
            continue
        chunk_start = start + len(raw_window) - len(raw_window.lstrip())
        chunk_end = end - (len(raw_window) - len(raw_window.rstrip()))

        while marker_index < len(markers) and markers[marker_index][0] < chunk_end:
            _, kind, title = markers[marker_index]
            active[kind] = title
            marker_index += 1

        locator: dict = {
            "char_start": chunk_start,
            "char_end": chunk_end,
            **active,
        }
        pages = sorted(
            {
                page_number
                for page_start, page_end, page_number in (page_spans or [])
                if chunk_start < page_end and chunk_end > page_start
            }
        )
        if pages:
            locator["page_start"] = pages[0]
            locator["page_end"] = pages[-1]
            locator["pages"] = pages
        records.append({"text": window, **locator})
        if end >= len(normalized):
            break
    return records


async def _stage(
    name: str,
    progress: ProgressCallback | None,
    should_cancel: CancelCheck | None,
) -> None:
    if should_cancel is not None and await should_cancel():
        raise IngestionCancelled("Ingestion cancelled")
    if progress is not None:
        await progress(name)


async def upload_document(
    filename: str,
    content: bytes,
    metadata: dict | None = None,
    doc_key: str | None = None,
    ocr: bool | None = None,
    progress: ProgressCallback | None = None,
    should_cancel: CancelCheck | None = None,
    **_,
) -> dict:
    """Convert to Markdown, chunk, embed and upsert one document into Qdrant.

    ``metadata`` (subject/grade/lang/doc_type/lab_id …) is merged into every
    chunk payload so retrieval can be filtered by lab context. ``doc_key`` is
    the stable identity for replace-semantics — defaults to ``filename`` for
    single uploads; bulk ingest passes the relative path to avoid collisions
    between same-named lab files.

    Returns ``{file_id, filename, status, chunks}``. ``status`` is ``"ready"``
    on success or ``"empty"`` if the document yielded no usable text. Raises
    ``ValueError`` for unsupported extensions; embedding / vector-store failures
    propagate as ``LLMError`` subclasses.
    """
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Supported: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    key = doc_key or filename
    doc_id = _doc_id(key)
    lang = (metadata or {}).get("lang")
    ocr_enabled = settings.OCR_ENABLED if ocr is None else ocr
    await _stage("extracting", progress, should_cancel)
    text = await asyncio.to_thread(
        to_markdown, filename, content, ocr=ocr_enabled, lang=lang
    )
    page_spans = (
        await asyncio.to_thread(_pdf_page_spans, content, text)
        if suffix == ".pdf"
        else []
    )
    chunk_records = _chunk_records(text, page_spans=page_spans)
    chunks = [record["text"] for record in chunk_records]
    if not chunk_records:
        logger.info(
            "Ingested '%s' -> doc_id=%s status=empty (existing chunks preserved)",
            key,
            doc_id,
        )
        return {"file_id": doc_id, "filename": filename, "status": "empty", "chunks": 0}

    await _stage("embedding", progress, should_cancel)
    embeddings_list = await embeddings.embed_texts(chunks)
    await _stage("indexing", progress, should_cancel)
    await vectorstore.ensure_collection()

    base_payload = {k: v for k, v in (metadata or {}).items() if v is not None}
    base_payload.setdefault("source_path", base_payload.get("source") or key)
    base_payload.setdefault(
        "source_type", base_payload.get("doc_type") or "document"
    )
    base_payload.setdefault("file_type", suffix.removeprefix("."))
    generation = uuid.uuid4().hex
    points = [
        {
            "id": uuid.uuid5(_DOC_NAMESPACE, f"{key}:{generation}:{i}").hex,
            "dense": emb.dense,
            "sparse_indices": emb.sparse_indices,
            "sparse_values": emb.sparse_values,
            "payload": {
                **base_payload,
                "doc_id": doc_id,
                "generation": generation,
                "filename": filename,
                "chunk_index": i,
                "chunk_count": len(chunk_records),
                **{k: v for k, v in record.items() if k != "text"},
                "text": record["text"],
            },
        }
        for i, (record, emb) in enumerate(zip(chunk_records, embeddings_list))
    ]

    n = await vectorstore.upsert_points(points)
    logger.info("Ingested '%s' -> doc_id=%s status=ready chunks=%d", key, doc_id, n)
    return {"file_id": doc_id, "filename": filename, "status": "ready", "chunks": n}


# ── Bulk corpus ingest + manifest ────────────────────────────────────────────


def iter_corpus_files(root: str) -> list[Path]:
    """All supported files under ``root`` (recursive), sorted for determinism."""
    base = Path(root)
    if not base.is_dir():
        raise ValueError(f"Corpus root is not a directory: {root}")
    return sorted(
        p
        for p in base.rglob("*")
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def _missing_metadata(meta: dict) -> list[str]:
    required = ["subject", "grade", "lang"]
    if meta.get("doc_type") == "lab_instruction":
        required.extend(("lab_number", "lab_id"))
    return [field for field in required if meta.get(field) is None]


def resolve_corpus_scope(root: str, subtree: str = "") -> tuple[Path, Path]:
    root_path = Path(root).resolve()
    if not root_path.is_dir():
        raise ValueError(f"Corpus root is not a directory: {root}")
    relative = Path(subtree or ".")
    if relative.is_absolute():
        raise ValueError("Corpus subtree must be relative to CORPUS_ROOT")
    scope = (root_path / relative).resolve()
    if not scope.is_relative_to(root_path):
        raise ValueError("Corpus subtree must remain inside CORPUS_ROOT")
    if not scope.is_dir():
        raise ValueError(f"Corpus subtree is not a directory: {subtree}")
    return root_path, scope


def scan_corpus_tree(
    root: str,
    *,
    subtree: str = "",
    only: str | None = None,
) -> dict:
    if subtree and only is not None:
        raise ValueError("subtree and only cannot be combined")
    root_path, scope = resolve_corpus_scope(root, subtree)
    files = iter_corpus_files(str(scope))
    present_doc_ids = {_doc_id(str(path.relative_to(root_path))) for path in files}
    candidates: list[dict] = []
    skipped: list[dict] = []
    errors: list[dict] = []
    filtered = 0
    for path in files:
        if only and only not in str(path):
            filtered += 1
            continue
        meta = corpus_meta.parse_path(str(path), corpus_root=str(root_path))
        if meta is None:
            skipped.append(
                {
                    "source": str(path.relative_to(root_path)),
                    "error": "Unrecognised corpus path",
                }
            )
            continue
        missing = _missing_metadata(meta)
        if missing:
            issue = {
                "source": meta["source"],
                "error": f"Incomplete metadata: {', '.join(missing)}",
            }
            skipped.append(issue)
            errors.append(issue)
            continue
        candidates.append(
            {
                "path": str(path),
                "metadata": meta,
                "doc_id": _doc_id(meta["source"]),
            }
        )

    lab_counts = Counter(
        item["metadata"]["lab_id"]
        for item in candidates
        if item["metadata"]["doc_type"] == "lab_instruction"
    )
    duplicate_lab_ids = {
        lab_id for lab_id, count in lab_counts.items() if count > 1
    }
    accepted = []
    for item in candidates:
        lab_id = item["metadata"].get("lab_id")
        if lab_id in duplicate_lab_ids:
            issue = {
                "source": item["metadata"]["source"],
                "error": f"Duplicate lab_id: {lab_id}",
            }
            skipped.append(issue)
            errors.append(issue)
            continue
        accepted.append(item)

    counts_by_type = Counter(item["metadata"]["doc_type"] for item in accepted)
    counts_by_language = Counter(item["metadata"]["lang"] for item in accepted)
    return {
        "root": str(root_path),
        "subtree": subtree,
        "total": len(files),
        "filtered": filtered,
        "candidates": accepted,
        "skipped": skipped,
        "errors": errors,
        "present_doc_ids": present_doc_ids,
        "duplicate_lab_ids": sorted(duplicate_lab_ids),
        "counts_by_type": dict(sorted(counts_by_type.items())),
        "counts_by_language": dict(sorted(counts_by_language.items())),
    }


async def prune_missing_corpus_documents(present_doc_ids: set[str]) -> int:
    pruned = 0
    for document in await vectorstore.list_documents():
        source = document.get("source_path")
        file_id = document.get("file_id")
        if (
            source
            and file_id
            and not source.startswith("admin_uploads/")
            and corpus_meta.parse_path(source) is not None
            and file_id not in present_doc_ids
            and await vectorstore.delete_document(file_id)
        ):
            pruned += 1
    return pruned


async def bulk_ingest_tree(
    root: str,
    *,
    ocr: bool | None = None,
    only: str | None = None,
    prune: bool = False,
) -> dict:
    """Walk ``root``, derive metadata from each path and ingest every file.

    Returns a summary
    ``{root, total, ready, empty, skipped, filtered, errors, documents}``.
    Files whose path is not under a recognised corpus tier are skipped (logged).
    ``only`` keeps the ingest root (so doc_ids stay stable) but restricts the run
    to files whose path contains that substring — used to OCR just one subtree
    (e.g. ``Биология/рус``) without re-embedding the whole corpus. ``ocr`` enables
    the scanned-document OCR fallback (see :func:`to_markdown`). Per-file failures
    are collected rather than aborting the whole run. ``prune`` explicitly removes
    stored corpus documents missing from this complete snapshot and cannot be
    combined with ``only``.
    """
    if prune and only is not None:
        raise ValueError("prune cannot be combined with a filtered bulk ingest")
    scan = scan_corpus_tree(root, only=only)
    ocr_enabled = settings.OCR_ENABLED if ocr is None else ocr
    summary = {
        "root": root,
        "total": scan["total"],
        "ready": 0,
        "empty": 0,
        "skipped": len(scan["skipped"]),
        "filtered": scan["filtered"],
        "errors": list(scan["errors"]),
        "documents": [],
        "documents_by_language": {},
        "pruned": 0,
    }
    for issue in scan["skipped"]:
        if issue in scan["errors"]:
            continue
        logger.info("Skip (unrecognised path): %s", issue["source"])
    for item in scan["candidates"]:
        path = Path(item["path"])
        meta = item["metadata"]
        try:
            result = await upload_document(
                path.name,
                path.read_bytes(),
                metadata=meta,
                doc_key=meta["source"],
                ocr=ocr_enabled,
            )
        except Exception as exc:  # noqa: BLE001 - keep going, record the failure
            summary["errors"].append({"source": meta["source"], "error": str(exc)})
            logger.warning("Ingest failed for %s: %s", path, exc)
            continue
        summary["ready" if result["status"] == "ready" else "empty"] += 1
        summary["documents"].append({**meta, **result})
        language = meta.get("lang")
        if language and result["status"] == "ready":
            summary["documents_by_language"][language] = (
                summary["documents_by_language"].get(language, 0) + 1
            )

    if prune and scan["candidates"] and not summary["errors"]:
        try:
            summary["pruned"] = await prune_missing_corpus_documents(
                scan["present_doc_ids"]
            )
        except Exception as exc:  # noqa: BLE001 - report prune failure with the run
            summary["errors"].append({"source": "<prune>", "error": str(exc)})
    logger.info(
        "Bulk ingest of %s: %d ready, %d empty, %d skipped, %d filtered, %d errors",
        root,
        summary["ready"],
        summary["empty"],
        summary["skipped"],
        summary["filtered"],
        len(summary["errors"]),
    )
    return summary


def build_manifest(root: str) -> dict:
    """Scan the corpus tree and report lab completeness (no embedding).

    Produces ``{labs: {lab_id: {...}}, missing_metadata: [...], textbooks: N}``.
    A lab is ``complete`` when its procedure file extracts enough text, ``stub``
    when present-but-thin, and we record which languages exist per
    subject/grade/number so a missing translation is visible. This is an offline
    report; the request path treats "no instruction in Qdrant" as incomplete.
    """
    labs: dict[str, dict] = {}
    duplicate_lab_ids: set[str] = set()
    missing_metadata: list[str] = []
    textbooks = 0
    textbooks_by_language: dict[str, int] = {}
    labs_by_language: dict[str, int] = {}

    for path in iter_corpus_files(root):
        meta = corpus_meta.parse_path(str(path), corpus_root=root)
        if meta is None:
            missing_metadata.append(str(path))
            continue
        missing = _missing_metadata(meta)
        if missing:
            missing_metadata.append(meta["source"])
            continue
        if meta["doc_type"] == "textbook":
            textbooks += 1
            language = meta.get("lang")
            if language:
                textbooks_by_language[language] = (
                    textbooks_by_language.get(language, 0) + 1
                )
            continue
        lab_id = meta.get("lab_id")
        if not lab_id:
            missing_metadata.append(meta["source"])
            continue
        if lab_id in duplicate_lab_ids:
            missing_metadata.append(meta["source"])
            continue
        if lab_id in labs:
            previous = labs.pop(lab_id)
            duplicate_lab_ids.add(lab_id)
            missing_metadata.extend((previous["source"], meta["source"]))
            continue
        try:
            text = to_markdown(path.name, path.read_bytes())
        except Exception:  # noqa: BLE001
            text = ""
        chars = len(re.sub(r"\s+", " ", text).strip())
        labs[lab_id] = {
            "lab_id": lab_id,
            "subject": meta["subject"],
            "grade": meta["grade"],
            "lang": meta["lang"],
            "lab_number": meta["lab_number"],
            "source": meta["source"],
            "chars": chars,
            "status": "complete" if chars >= _STUB_CHARS else "stub",
        }
    for lab in labs.values():
        language = lab.get("lang")
        if language:
            labs_by_language[language] = labs_by_language.get(language, 0) + 1

    return {
        "labs": dict(sorted(labs.items())),
        "missing_metadata": sorted(missing_metadata),
        "textbooks": textbooks,
        "textbooks_by_language": dict(sorted(textbooks_by_language.items())),
        "labs_by_language": dict(sorted(labs_by_language.items())),
    }


def write_manifest(root: str, out_path: str) -> dict:
    """Build the manifest for ``root`` and write it to ``out_path`` as JSON."""
    manifest = build_manifest(root)
    Path(out_path).write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return manifest


async def list_documents(**_) -> list[dict]:
    """List ingested documents (one entry per ``doc_id``)."""
    return await vectorstore.list_documents()


async def delete_document(file_id: str, **_) -> bool:
    """Delete a document's chunks; return False if it did not exist."""
    return await vectorstore.delete_document(file_id)


async def corpus_status(**_) -> dict:
    """Return the Qdrant collection's status, point count and document count."""
    return await vectorstore.collection_status()
